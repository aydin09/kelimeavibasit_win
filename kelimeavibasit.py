from tkinter import *
import random
import os
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.shared import Mm, Cm
from docx.shared import Length

pencere = Tk()
pencere.tk_setPalette("light blue")
pencere.resizable(width=FALSE ,height=FALSE)
pencere.geometry("400x200+300+100")
baslik = pencere.title("KELİME AVI")
img=PhotoImage(file='kelime.png')
pencere.tk.call('wm','iconphoto',pencere._w,img)

klm=[]
klm1=[]
klm2=[]
k1=[]
k2=[]
k3=[]
k4=[]
k5=[]
k6=[]
k7=[]
k8=[]
k9=[]
k10=[]
k11=[]
k12=[]
k13=[]

alfabe=["C","D","F","Ğ","H","L","Ş","Z"]
abc1=[]
abc2=[]
abc3=[]
abc4=[]
abc5=[]
abc6=[]
abc7=[]
abc8=[]
abc9=[]
abc10=[]
abc11=[]
abc12=[]
abc13=[]

kelime = ["LALE","KEK","ELEK","KEKLİK","KEKİK","NANE","NAL","ENİK","KONAK","KOLA","KANO","KALEM","EKMEK","LİMON","KİLİM","KALEMLİK","LOKUM","UN",
          "KUM","KUKLA","LİMONLUK","KULAKLIK","TEL","TAKA","ATLET","TEKE","KOT","TEKNE","KUTU","AT","ET","OT","OLTA","TOKA","KİLİT","ETİKET",
          "KOLTUK","ETEK","ÜTÜ","KÜLLÜK","KÜTÜK","AKÜ","NAYLON","LEYLAK","KAYAK","YELKEN","KOLYE","YELEK","AYNA","TAY","KAYKAY","YAY","KİMYON","ÖNLÜK",
          "NAR","KEMER","ARMUT","TEKERLEK","YUMURTA","KÖMÜR","ATARİ","ERİK","TRAKTÖR","RAKET","TERLİK","AYRAN","TARLA","KARTON","TANKER","TARAK","TERE",
          "ROKA","KÜREK","KORNA","KARYOLA","ATKI","ALTIN","TAKI","TIRMIK","TARTI","KINA","KAYIK","KULAKLIK","TIR","DUT","DÜDÜK","DÖNER",
          "DANA","DANTEL","DEMİR","DEMLİK","ÖRDEK","KARDELEN","ORKİDE","DARI","TAKSİ",
          "KESTANE","KAYISI","TOST","ASKI","SODA","SANDALYE","SAKSI","SANDAL","LASTİK","SOSİS","YASTIK","KASA","SANDIK","SARAY","TESTERE","BALIK","BALON",
          "BAYRAK","BÖREK","AYAKKABI","SABUN","BİLET","LEBLEBİ","BALTA","BANT","BİBER","BOT","BAKLA","BADEM","BASTON","KABAK","BARBUNYA","BATTANİYE","ARABA",
          "BONE","TABAK","BİBLO","TABLO","SÜMBÜL","DÜRBÜN","KİBRİT","TABLET","BASKÜL","ABAKÜS","ELBİSE","TABURE","BİBERON","TORBA","ROBOT",
          "BİLYE","RADYO","ZİL","BİLEZİK","ZEYTİN","YÜZÜK","ÜZÜM","KUZU","EMZİK","TUZ","SÖZLÜK","ZIMBA","TERAZİ","BUZDOLABI","BEZ","ROZET","SAKIZ",
          "ZURNA","YAZLIK","BALYOZ","BENZİN","BEZELYE","ZERDALİ","ÇİLEK","REÇEL","ÇÖREK","ÇAYDANLIK","ÇİKOLATA","ÇAY","ÇADIR","KEÇİ","ÇANTA","ÇEKİÇ","SALÇA",
          "ÇİZME","ÇİMENTO","ÇÖMLEK","ÇANAK","SÜTLAÇ","TARÇIN","ÇİÇEKLİK","KILIÇ","KEÇE","ÇÖKELEK","DERGİ","GİYSİ","SİLGİ","GÖMLEK",
          "SÜNGER","YORGAN","BULGUR","BİLGİSAYAR","SÜZGEÇ","GÖZLEME","GÖZLÜK","GIDA","ŞAL","KUŞ","ATAŞ","ŞORT","TAŞIT","ŞEKER","BEŞİK","TURŞU","ÇEŞME",
          "TEBEŞİR","ŞEMSİYE","OCAK","KANCA","CEKET","SUCUK","ZİNCİR","BONCUK","CÜZDAN","OYUNCAK","TENCERE","KÜP","PİL","DOLAP","KUPA",
          "POŞET","PİDE","ÇAPA","KÜPE","PASTA","PANO","PERDE","ŞAPKA","KANEPE","PUSULA","PANTOLON","RAPTİYE","SÜPÜRGE","PİYANO","KEBAP","PATETES","PATLICAN",
          "TURP","ISPANAK","PAPATYA","PATİK","PEKMEZ","PORTAKAL","PIRASA","KARPUZ","PEYNİR","ÇORAP","KİTAP","PENSE","PERGEL","PEÇETE","PİRİNÇ","PATEN","PALTO",
          "PESTİL","YELPAZE","HALI","HOROZ","HEDİYE","HIRKA","HORTUM","HİNDİ","LAHANA","BAHARAT","HUNİ","HALAT","SÜRAHİ","VİDA","KOVA","AYVA","BAVUL","HAVLU",
          "KAVANOZ","NEVRESİM","VALİZ","VİŞNE","BAKLAVA","KAVUN","CEVİZ","REVANİ","KİVİ","VAZO","TAVA","TAVUK","CİVCİV","KAVAL","DAVUL","HAVUÇ","ÇİVİ",
          "ÇUVAL","KAHVE","BİSKÜVİ","ELDİVEN","TORNAVİDA","TELEVİZYON","KEREVİZ","CETVEL","AVİZE","KAĞIT","POĞAÇA","YAĞ","OĞLAK","LEĞEN","TIĞ",
          "SOĞAN","TUĞLA","PAPAĞAN","ZEYTİNYAĞI","BUĞDAY","YAĞMURLUK","FES","FAN","FAR","FİŞ","ÇARŞAF","FİDAN","FİLE","KAFES","DEFTER","FENER","FIRIN",
          "FISTIK","GOFRET","KÖFTE","FASULYE","KADİFE","KIYAFET","TELEFON","KADAYIF","KARANFİL","FİNCAN","FIRÇA","YUFKA","FİDE","FARAŞ","KÜNEFE",
          "JETON","ABAJUR","DETERJAN","JÖLE","BANDAJ","RUJ","JİLE","AJANDA","OJE","JELİBON","JENERATÖR","JANT","BUJİ","KAJU","JALUZİ","ELA","LALE","ALİ",
          "AKİLE","NİL","NAİL","İNAN","NAİLE","NALAN","ALKAN","KENAN","KAAN","EKİN","OKAN","EMİN","EMİNE","EMEL","ALİM",
          "ATA","TAN","TALAT","TEKİN","ATAKAN","ONAT","TEMEL","ALTAN","ATİLLA","UTKU","TUNA","NİMET","ÜMİT","ÜLKÜ",
          "TÜLİN","ÜNAL","KAYA","LEYLA","OYA","EYLÜL","AYLA","OKTAY","AYTEN","EYMEN","KUTAY","AYLİN","AYKUT","TÜLAY","ÖYKÜ","ÖNAL","ÖKTEM","RANA","ONUR",
          "ÖMER","TURAN","EREN","ERTAN","NUR","NURİ","ERAY","ÖMÜR","EMRE","NURAN","ERKAN","EROL","AYNUR","TANER","ITIR","ANIL","TARIK",
          "AKIN","ALKIN","DEMET","DİLEK","DİDEM","EDA","KADİR","ARDA","DİLARA","DERYA","AYDIN","DURU","ADEM","ERDAL","ÖNDER","ENDER","ERDEM","SEDA","SELİN",
          "SELİM","SİNAN","SADIK","SEREN","ESİN","ASYA","ERSİN","SUAT","SITKI","SONER","DURSUN","SİNEM","YASİN","SÜMER","ASLI","ESRA",
          "SEMRA","SILA","SELİNAY","KASIM","ASU","SUDE","BANU","BETÜL","BURAK","BİROL","BEKİR","SUDE","EBRU","BERAT","SABRİ","BERK","BERİL","BERKAY",
          "BASRİ","BARAN","BERNA","BİRKAN","BÜLENT","BORA","BUKET","BİLAL","KÜBRA","SİBEL","BESTE","BERKE","KUBAT","TİBET","BERKANT","BAKİ","ÖZNUR","ZEKİ","YILDIZ",
          "KEZBAN","NAZLI","ÖZEN","OZAN","ÖZLEM","DENİZ","RIZA","ZERRİN","ZEKİYE","İZZET","REMZİYE","REMZİ","KAZIM","AZİZ","SUZAN","BEYZA","ARZU","ÖZKAN",
          "ZÜBEYDE","SEZEN","AZRA","NİYAZİ","NAZMİ","ÖZER","NAZMİYE","ZİYA","SEZER","ZEYNEL","ZÜMRÜT","ÇINAR","ÇETİN","ERDİNÇ","SEÇİL","AYÇA","ÇELİK",
          "TUNÇ","SEÇKİN","SERTAÇ","AÇELYA","YALÇIN","ÇİLEM","EGE","EZGİ","ÖZGE","ENGİN","BİLGE","DUYGU","TOLGA","BELGİN","BENGÜ","NURGÜL","ÖZGÜR",
          "SEZGİN","ŞULE","NEŞE","ŞAKİR","BARIŞ","ŞERMİN","AYŞE","YAŞAR","YEŞİM","ŞENAY","ŞENER","ŞÜKRÜ","ŞÜKRAN","ŞEBNEM","ŞENGÜL","ŞADİYE","AYŞEGÜL","BÜŞRA",
          "RÜŞTÜ","CEM","CENK","CAN","ECE","NACİ","NACİYE","ECEM","CEREN","CEMİL","CANAN","CANER","BURCU","CEMRE","CELAL","SACİT","CENGİZ","CANDAN","CANSU",
          "CEZMİ","CEYDA","CEYLİN","CANBERK","CÜNEYT","NECATİ","CENNET","ECRİN","İPEK","SARP","ALPER","SERAP","SERPİL","RECEP","ZEYNEP","ALP","ALPAY","HİLMİ",
          "HANDE","HANDAN","HÜLYA","HASAN","SEHER","HARUN","HAKAN","SEMİH","ŞAHİN","BAHRİ","HİLAL","REHA","HAYRİYE","HAYRİ","HATİCE","HAKKI",
          "HURİYE","NİHAL","HALİT","HALE","SALİH","NUH","HAVVA","HÜSEYİN","TAHA","LEVENT","NEVİN","CEVDET","YAVUZ","AVNİ","SEVGİ","TUĞÇE","ÇİĞDEM","BUĞRA",
          "ÇAĞLAR","UĞUR","DOĞUŞ","DOĞAN","TUĞBA","AYTUĞ","DOĞA","TUĞRUL","UFUK","FİLİZ","ZAFER","FERİT","FUNDA","FUAT","FATMA","FEYZA","FİKRET","FARUK","EFE",
          "FAHRİ","FIRAT","FURKAN","ELİF","YUSUF","SEFA","FADİME","FATİH","FERDİ","FERİDE","FULYA","FAİK","FAZIL","ŞAFAK","NEJAT","TANJU"]

def kelime_avi():
    pencere.destroy()
    for k in random.sample(kelime,13):
        klm.append(k)
        klm1.append(k)

    for h in range(len(klm1)):
        for p in random.sample(klm1,1):
            klm2.append(p)
            klm1.remove(p)
        
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(20)
    font.bold = True

    footer_section = document.sections[0]
    footer = footer_section.footer

    footer_text = footer.paragraphs[0]
    footer_text.text = "\t\twww.egitimhane.com"

    paragraph = document.add_paragraph("KELİME AVI")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = document.add_table(rows=1, cols=1,style = 'Table Grid')
    cell = table.cell(0,0)
    table.cell(0,0).paragraphs[0].add_run("Adı Soyadı:")

    paragraph = document.add_paragraph()

    paragraph = document.add_paragraph(klm[0]+"  "+klm[1]+"  "+klm[2]+"  "+klm[3]+"  "+klm[4]+"  "+klm[5]+"  "+klm[6]+"  "+klm[7]+
                                       "  "+klm[8]+"  "+klm[9]+"  "+klm[10]+"  "+klm[11]+"  "+klm[12])  
    paragraph = document.add_paragraph()

    table = document.add_table(rows=13, cols=10,style = 'Table Grid')
    cell = table.cell(0,0)
    
    for a in random.sample(alfabe,8):
        abc1.append(a)

    for a in random.sample(alfabe,8):
        abc2.append(a)

    for a in random.sample(alfabe,8):
        abc3.append(a)

    for a in random.sample(alfabe,8):
        abc4.append(a)

    for a in random.sample(alfabe,8):
        abc5.append(a)

    for a in random.sample(alfabe,8):
        abc6.append(a)

    for a in random.sample(alfabe,8):
        abc7.append(a)

    for a in random.sample(alfabe,8):
        abc8.append(a)

    for a in random.sample(alfabe,8):
        abc9.append(a)

    for a in random.sample(alfabe,8):
        abc10.append(a)

    for a in random.sample(alfabe,8):
        abc11.append(a)

    for a in random.sample(alfabe,8):
        abc12.append(a)

    for a in random.sample(alfabe,8):
        abc13.append(a)    
    
    if len(klm2[0]) == 10:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,3).paragraphs[0].add_run(k1[3])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,4).paragraphs[0].add_run(k1[4])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(k1[5])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(k1[6])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(k1[7])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(k1[8])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(k1[9])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[0]) == 9:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,3).paragraphs[0].add_run(k1[3])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,4).paragraphs[0].add_run(k1[4])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(k1[5])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(k1[6])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(k1[7])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(k1[8])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[0])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
  
    if len(klm2[0]) == 8:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,3).paragraphs[0].add_run(k1[3])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,4).paragraphs[0].add_run(k1[4])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(k1[5])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(k1[6])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(k1[7])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(abc1[0])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[1])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[0]) == 7:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,3).paragraphs[0].add_run(k1[3])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,4).paragraphs[0].add_run(k1[4])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(k1[5])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(k1[6])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(abc1[0])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(abc1[1])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[2])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[0]) == 6:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,3).paragraphs[0].add_run(k1[3])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,4).paragraphs[0].add_run(k1[4])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(k1[5])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(abc1[0])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(abc1[1])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(abc1[2])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[3])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[0]) == 5:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,3).paragraphs[0].add_run(k1[3])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,4).paragraphs[0].add_run(k1[4])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(abc1[0])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(abc1[1])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(abc1[2])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(abc1[3])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[4])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[0]) == 4:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,3).paragraphs[0].add_run(k1[3])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,4).paragraphs[0].add_run(abc1[0])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(abc1[1])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(abc1[2])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(abc1[3])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(abc1[4])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[5])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER        

    if len(klm2[0]) == 3:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,2).paragraphs[0].add_run(k1[2])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,3).paragraphs[0].add_run(abc1[0])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,4).paragraphs[0].add_run(abc1[1])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(abc1[2])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(abc1[3])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(abc1[4])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(abc1[5])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[6])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[0]) == 2:
        for i in klm2[0]:
            k1.append(i)

        table.cell(0,0).paragraphs[0].add_run(k1[0])
        table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(0,1).paragraphs[0].add_run(k1[1])
        table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,2).paragraphs[0].add_run(abc1[0])
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,3).paragraphs[0].add_run(abc1[1])
        table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,4).paragraphs[0].add_run(abc1[2])
        table.cell(0,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,5).paragraphs[0].add_run(abc1[3])
        table.cell(0,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,6).paragraphs[0].add_run(abc1[4])
        table.cell(0,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,7).paragraphs[0].add_run(abc1[5])
        table.cell(0,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,8).paragraphs[0].add_run(abc1[6])
        table.cell(0,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(0,9).paragraphs[0].add_run(abc1[7])
        table.cell(0,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[1]) == 10:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,3).paragraphs[0].add_run(k2[3])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,4).paragraphs[0].add_run(k2[4])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(k2[5])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(k2[6])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(k2[7])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(k2[8])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(k2[9])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[1]) == 9:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,3).paragraphs[0].add_run(k2[3])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,4).paragraphs[0].add_run(k2[4])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(k2[5])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(k2[6])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(k2[7])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(k2[8])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[0])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[1]) == 8:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,3).paragraphs[0].add_run(k2[3])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,4).paragraphs[0].add_run(k2[4])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(k2[5])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(k2[6])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(k2[7])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(abc2[0])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[1])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[1]) == 7:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,3).paragraphs[0].add_run(k2[3])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,4).paragraphs[0].add_run(k2[4])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(k2[5])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(k2[6])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(abc2[0])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(abc2[1])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[2])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[1]) == 6:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,3).paragraphs[0].add_run(k2[3])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,4).paragraphs[0].add_run(k2[4])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(k2[5])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(abc2[0])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(abc2[1])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(abc2[2])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[3])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[1]) == 5:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,3).paragraphs[0].add_run(k2[3])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,4).paragraphs[0].add_run(k2[4])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(abc2[0])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(abc2[1])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(abc2[2])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(abc2[3])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[4])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[1]) == 4:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,3).paragraphs[0].add_run(k2[3])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,4).paragraphs[0].add_run(abc2[0])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(abc2[1])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(abc2[2])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(abc2[3])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(abc2[4])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[5])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[1]) == 3:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,2).paragraphs[0].add_run(k2[2])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,3).paragraphs[0].add_run(abc2[0])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,4).paragraphs[0].add_run(abc2[1])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(abc2[2])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(abc2[3])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(abc2[4])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(abc2[5])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[6])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[1]) == 2:
        for i in klm2[1]:
            k2.append(i)

        table.cell(1,0).paragraphs[0].add_run(k2[0])
        table.cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(1,1).paragraphs[0].add_run(k2[1])
        table.cell(1,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,2).paragraphs[0].add_run(abc2[0])
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,3).paragraphs[0].add_run(abc2[1])
        table.cell(1,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,4).paragraphs[0].add_run(abc2[2])
        table.cell(1,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,5).paragraphs[0].add_run(abc2[3])
        table.cell(1,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,6).paragraphs[0].add_run(abc2[4])
        table.cell(1,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,7).paragraphs[0].add_run(abc2[5])
        table.cell(1,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,8).paragraphs[0].add_run(abc2[6])
        table.cell(1,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1,9).paragraphs[0].add_run(abc2[7])
        table.cell(1,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[2]) == 10:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,3).paragraphs[0].add_run(k3[3])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,4).paragraphs[0].add_run(k3[4])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(k3[5])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(k3[6])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(k3[7])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(k3[8])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(k3[9])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[2]) == 9:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,3).paragraphs[0].add_run(k3[3])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,4).paragraphs[0].add_run(k3[4])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(k3[5])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(k3[6])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(k3[7])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(k3[8])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[0])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[2]) == 8:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,3).paragraphs[0].add_run(k3[3])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,4).paragraphs[0].add_run(k3[4])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(k3[5])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(k3[6])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(k3[7])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(abc3[0])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[1])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[2]) == 7:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,3).paragraphs[0].add_run(k3[3])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,4).paragraphs[0].add_run(k3[4])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(k3[5])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(k3[6])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(abc3[0])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(abc3[1])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[2])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[2]) == 6:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,3).paragraphs[0].add_run(k3[3])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,4).paragraphs[0].add_run(k3[4])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(k3[5])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(abc3[0])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(abc3[1])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(abc3[2])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[3])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[2]) == 5:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,3).paragraphs[0].add_run(k3[3])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,4).paragraphs[0].add_run(k3[4])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(abc3[0])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(abc3[1])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(abc3[2])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(abc3[3])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[4])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[2]) == 4:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,3).paragraphs[0].add_run(k3[3])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,4).paragraphs[0].add_run(abc3[0])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(abc3[1])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(abc3[2])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(abc3[3])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(abc3[4])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[5])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[2]) == 3:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,2).paragraphs[0].add_run(k3[2])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,3).paragraphs[0].add_run(abc3[0])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,4).paragraphs[0].add_run(abc3[1])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(abc3[2])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(abc3[3])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(abc3[4])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(abc3[5])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[6])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[2]) == 2:
        for i in klm2[2]:
            k3.append(i)

        table.cell(2,0).paragraphs[0].add_run(k3[0])
        table.cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(2,1).paragraphs[0].add_run(k3[1])
        table.cell(2,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,2).paragraphs[0].add_run(abc3[0])
        table.cell(2,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,3).paragraphs[0].add_run(abc3[1])
        table.cell(2,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,4).paragraphs[0].add_run(abc3[2])
        table.cell(2,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,5).paragraphs[0].add_run(abc3[3])
        table.cell(2,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,6).paragraphs[0].add_run(abc3[4])
        table.cell(2,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,7).paragraphs[0].add_run(abc3[5])
        table.cell(2,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,8).paragraphs[0].add_run(abc3[6])
        table.cell(2,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(2,9).paragraphs[0].add_run(abc3[7])
        table.cell(2,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[3]) == 10:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,3).paragraphs[0].add_run(k4[3])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,4).paragraphs[0].add_run(k4[4])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(k4[5])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(k4[6])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(k4[7])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(k4[8])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(k4[9])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[3]) == 9:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,3).paragraphs[0].add_run(k4[3])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,4).paragraphs[0].add_run(k4[4])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(k4[5])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(k4[6])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(k4[7])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(k4[8])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[0])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[3]) == 8:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,3).paragraphs[0].add_run(k4[3])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,4).paragraphs[0].add_run(k4[4])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(k4[5])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(k4[6])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(k4[7])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(abc4[0])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[1])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[3]) == 7:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,3).paragraphs[0].add_run(k4[3])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,4).paragraphs[0].add_run(k4[4])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(k4[5])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(k4[6])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(abc4[0])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(abc4[1])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[2])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[3]) == 6:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,3).paragraphs[0].add_run(k4[3])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,4).paragraphs[0].add_run(k4[4])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(k4[5])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(abc4[0])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(abc4[1])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(abc4[2])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[3])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[3]) == 5:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,3).paragraphs[0].add_run(k4[3])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,4).paragraphs[0].add_run(k4[4])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(abc4[0])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(abc4[1])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(abc4[2])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(abc4[3])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[4])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[3]) == 4:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,3).paragraphs[0].add_run(k4[3])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,4).paragraphs[0].add_run(abc4[0])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(abc4[1])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(abc4[2])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(abc4[3])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(abc4[4])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[5])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[3]) == 3:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,2).paragraphs[0].add_run(k4[2])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,3).paragraphs[0].add_run(abc4[0])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,4).paragraphs[0].add_run(abc4[1])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(abc4[2])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(abc4[3])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(abc4[4])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(abc4[5])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[6])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[3]) == 2:
        for i in klm2[3]:
            k4.append(i)

        table.cell(3,0).paragraphs[0].add_run(k4[0])
        table.cell(3,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(3,1).paragraphs[0].add_run(k4[1])
        table.cell(3,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,2).paragraphs[0].add_run(abc4[0])
        table.cell(3,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,3).paragraphs[0].add_run(abc4[1])
        table.cell(3,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,4).paragraphs[0].add_run(abc4[2])
        table.cell(3,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,5).paragraphs[0].add_run(abc4[3])
        table.cell(3,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,6).paragraphs[0].add_run(abc4[4])
        table.cell(3,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,7).paragraphs[0].add_run(abc4[5])
        table.cell(3,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,8).paragraphs[0].add_run(abc4[6])
        table.cell(3,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(3,9).paragraphs[0].add_run(abc4[7])
        table.cell(3,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[4]) == 10:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,3).paragraphs[0].add_run(k5[3])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,4).paragraphs[0].add_run(k5[4])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(k5[5])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(k5[6])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(k5[7])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(k5[8])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(k5[9])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[4]) == 9:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,3).paragraphs[0].add_run(k5[3])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,4).paragraphs[0].add_run(k5[4])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(k5[5])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(k5[6])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(k5[7])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(k5[8])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[0])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[4]) == 8:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,3).paragraphs[0].add_run(k5[3])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,4).paragraphs[0].add_run(k5[4])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(k5[5])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(k5[6])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(k5[7])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(abc5[0])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[1])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[4]) == 7:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,3).paragraphs[0].add_run(k5[3])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,4).paragraphs[0].add_run(k5[4])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(k5[5])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(k5[6])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(abc5[0])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(abc5[1])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[2])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[4]) == 6:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,3).paragraphs[0].add_run(k5[3])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,4).paragraphs[0].add_run(k5[4])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(k5[5])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(abc5[0])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(abc5[1])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(abc5[2])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[3])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[4]) == 5:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,3).paragraphs[0].add_run(k5[3])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,4).paragraphs[0].add_run(k5[4])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(abc5[0])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(abc5[1])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(abc5[2])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(abc5[3])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[4])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[4]) == 4:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,3).paragraphs[0].add_run(k5[3])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,4).paragraphs[0].add_run(abc5[0])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(abc5[1])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(abc5[2])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(abc5[3])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(abc5[4])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[5])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[4]) == 3:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,2).paragraphs[0].add_run(k5[2])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,3).paragraphs[0].add_run(abc5[0])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,4).paragraphs[0].add_run(abc5[1])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(abc5[2])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(abc5[3])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(abc5[4])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(abc5[5])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[6])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[4]) == 2:
        for i in klm2[4]:
            k5.append(i)

        table.cell(4,0).paragraphs[0].add_run(k5[0])
        table.cell(4,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(4,1).paragraphs[0].add_run(k5[1])
        table.cell(4,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,2).paragraphs[0].add_run(abc5[0])
        table.cell(4,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,3).paragraphs[0].add_run(abc5[1])
        table.cell(4,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,4).paragraphs[0].add_run(abc5[2])
        table.cell(4,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,5).paragraphs[0].add_run(abc5[3])
        table.cell(4,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,6).paragraphs[0].add_run(abc5[4])
        table.cell(4,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,7).paragraphs[0].add_run(abc5[5])
        table.cell(4,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,8).paragraphs[0].add_run(abc5[6])
        table.cell(4,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(4,9).paragraphs[0].add_run(abc5[7])
        table.cell(4,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[5]) == 10:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,3).paragraphs[0].add_run(k6[3])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,4).paragraphs[0].add_run(k6[4])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(k6[5])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(k6[6])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(k6[7])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(k6[8])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(k6[9])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[5]) == 9:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,3).paragraphs[0].add_run(k6[3])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,4).paragraphs[0].add_run(k6[4])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(k6[5])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(k6[6])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(k6[7])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(k6[8])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[0])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[5]) == 8:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,3).paragraphs[0].add_run(k6[3])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,4).paragraphs[0].add_run(k6[4])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(k6[5])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(k6[6])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(k6[7])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(abc6[0])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[1])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[5]) == 7:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,3).paragraphs[0].add_run(k6[3])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,4).paragraphs[0].add_run(k6[4])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(k6[5])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(k6[6])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(abc6[0])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(abc6[1])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[2])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[5]) == 6:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,3).paragraphs[0].add_run(k6[3])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,4).paragraphs[0].add_run(k6[4])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(k6[5])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(abc6[0])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(abc6[1])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(abc6[2])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[3])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[5]) == 5:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,3).paragraphs[0].add_run(k6[3])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,4).paragraphs[0].add_run(k6[4])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(abc6[0])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(abc6[1])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(abc6[2])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(abc6[3])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[4])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[5]) == 4:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,3).paragraphs[0].add_run(k6[3])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,4).paragraphs[0].add_run(abc6[0])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(abc6[1])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(abc6[2])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(abc6[3])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(abc6[4])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[5])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[5]) == 3:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,2).paragraphs[0].add_run(k6[2])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,3).paragraphs[0].add_run(abc6[0])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,4).paragraphs[0].add_run(abc6[1])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(abc6[2])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(abc6[3])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(abc6[4])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(abc6[5])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[6])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[5]) == 2:
        for i in klm2[5]:
            k6.append(i)

        table.cell(5,0).paragraphs[0].add_run(k6[0])
        table.cell(5,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(5,1).paragraphs[0].add_run(k6[1])
        table.cell(5,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,2).paragraphs[0].add_run(abc6[0])
        table.cell(5,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,3).paragraphs[0].add_run(abc6[1])
        table.cell(5,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,4).paragraphs[0].add_run(abc6[2])
        table.cell(5,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,5).paragraphs[0].add_run(abc6[3])
        table.cell(5,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,6).paragraphs[0].add_run(abc6[4])
        table.cell(5,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,7).paragraphs[0].add_run(abc6[5])
        table.cell(5,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,8).paragraphs[0].add_run(abc6[6])
        table.cell(5,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(5,9).paragraphs[0].add_run(abc6[7])
        table.cell(5,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[6]) == 10:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,3).paragraphs[0].add_run(k7[3])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,4).paragraphs[0].add_run(k7[4])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(k7[5])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(k7[6])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(k7[7])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(k7[8])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(k7[9])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[6]) == 9:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,3).paragraphs[0].add_run(k7[3])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,4).paragraphs[0].add_run(k7[4])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(k7[5])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(k7[6])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(k7[7])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(k7[8])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[0])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[6]) == 8:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,3).paragraphs[0].add_run(k7[3])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,4).paragraphs[0].add_run(k7[4])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(k7[5])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(k7[6])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(k7[7])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(abc7[0])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[1])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[6]) == 7:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,3).paragraphs[0].add_run(k7[3])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,4).paragraphs[0].add_run(k7[4])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(k7[5])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(k7[6])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(abc7[0])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(abc7[1])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[2])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[6]) == 6:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,3).paragraphs[0].add_run(k7[3])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,4).paragraphs[0].add_run(k7[4])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(k7[5])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(abc7[0])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(abc7[1])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(abc7[2])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[3])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[6]) == 5:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,3).paragraphs[0].add_run(k7[3])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,4).paragraphs[0].add_run(k7[4])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(abc7[0])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(abc7[1])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(abc7[2])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(abc7[3])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[4])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[6]) == 4:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,3).paragraphs[0].add_run(k7[3])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,4).paragraphs[0].add_run(abc7[0])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(abc7[1])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(abc7[2])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(abc7[3])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(abc7[4])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[5])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[6]) == 3:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,2).paragraphs[0].add_run(k7[2])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,3).paragraphs[0].add_run(abc7[0])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,4).paragraphs[0].add_run(abc7[1])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(abc7[2])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(abc7[3])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(abc7[4])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(abc7[5])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[6])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[6]) == 2:
        for i in klm2[6]:
            k7.append(i)

        table.cell(6,0).paragraphs[0].add_run(k7[0])
        table.cell(6,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(6,1).paragraphs[0].add_run(k7[1])
        table.cell(6,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,2).paragraphs[0].add_run(abc7[0])
        table.cell(6,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,3).paragraphs[0].add_run(abc7[1])
        table.cell(6,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,4).paragraphs[0].add_run(abc7[2])
        table.cell(6,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,5).paragraphs[0].add_run(abc7[3])
        table.cell(6,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,6).paragraphs[0].add_run(abc7[4])
        table.cell(6,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,7).paragraphs[0].add_run(abc7[5])
        table.cell(6,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,8).paragraphs[0].add_run(abc7[6])
        table.cell(6,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(6,9).paragraphs[0].add_run(abc7[7])
        table.cell(6,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[7]) == 10:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,3).paragraphs[0].add_run(k8[3])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,4).paragraphs[0].add_run(k8[4])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(k8[5])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(k8[6])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(k8[7])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(k8[8])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(k8[9])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[7]) == 9:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,3).paragraphs[0].add_run(k8[3])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,4).paragraphs[0].add_run(k8[4])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(k8[5])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(k8[6])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(k8[7])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(k8[8])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[0])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[7]) == 8:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,3).paragraphs[0].add_run(k8[3])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,4).paragraphs[0].add_run(k8[4])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(k8[5])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(k8[6])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(k8[7])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(abc8[0])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[1])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[7]) == 7:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,3).paragraphs[0].add_run(k8[3])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,4).paragraphs[0].add_run(k8[4])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(k8[5])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(k8[6])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(abc8[0])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(abc8[1])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[2])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[7]) == 6:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,3).paragraphs[0].add_run(k8[3])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,4).paragraphs[0].add_run(k8[4])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(k8[5])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(abc8[0])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(abc8[1])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(abc8[2])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[3])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[7]) == 5:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,3).paragraphs[0].add_run(k8[3])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,4).paragraphs[0].add_run(k8[4])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(abc8[0])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(abc8[1])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(abc8[2])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(abc8[3])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[4])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[7]) == 4:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,3).paragraphs[0].add_run(k8[3])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,4).paragraphs[0].add_run(abc8[0])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(abc8[1])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(abc8[2])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(abc8[3])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(abc8[4])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[5])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[7]) == 3:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,2).paragraphs[0].add_run(k8[2])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,3).paragraphs[0].add_run(abc8[0])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,4).paragraphs[0].add_run(abc8[1])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(abc8[2])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(abc8[3])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(abc8[4])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(abc8[5])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[6])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[7]) == 2:
        for i in klm2[7]:
            k8.append(i)

        table.cell(7,0).paragraphs[0].add_run(k8[0])
        table.cell(7,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(7,1).paragraphs[0].add_run(k8[1])
        table.cell(7,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,2).paragraphs[0].add_run(abc8[0])
        table.cell(7,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,3).paragraphs[0].add_run(abc8[1])
        table.cell(7,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,4).paragraphs[0].add_run(abc8[2])
        table.cell(7,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,5).paragraphs[0].add_run(abc8[3])
        table.cell(7,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,6).paragraphs[0].add_run(abc8[4])
        table.cell(7,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,7).paragraphs[0].add_run(abc8[5])
        table.cell(7,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,8).paragraphs[0].add_run(abc8[6])
        table.cell(7,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(7,9).paragraphs[0].add_run(abc8[7])
        table.cell(7,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[8]) == 10:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,3).paragraphs[0].add_run(k9[3])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,4).paragraphs[0].add_run(k9[4])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(k9[5])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(k9[6])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(k9[7])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(k9[8])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(k9[9])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[8]) == 9:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,3).paragraphs[0].add_run(k9[3])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,4).paragraphs[0].add_run(k9[4])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(k9[5])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(k9[6])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(k9[7])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(k9[8])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[0])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[8]) == 8:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,3).paragraphs[0].add_run(k9[3])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,4).paragraphs[0].add_run(k9[4])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(k9[5])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(k9[6])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(k9[7])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(abc9[0])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[1])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[8]) == 7:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,3).paragraphs[0].add_run(k9[3])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,4).paragraphs[0].add_run(k9[4])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(k9[5])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(k9[6])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(abc9[0])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(abc9[1])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[2])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[8]) == 6:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,3).paragraphs[0].add_run(k9[3])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,4).paragraphs[0].add_run(k9[4])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(k9[5])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(abc9[0])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(abc9[1])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(abc9[2])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[3])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[8]) == 5:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,3).paragraphs[0].add_run(k9[3])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,4).paragraphs[0].add_run(k9[4])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(abc9[0])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(abc9[1])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(abc9[2])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(abc9[3])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[4])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[8]) == 4:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,3).paragraphs[0].add_run(k9[3])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,4).paragraphs[0].add_run(abc9[0])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(abc9[1])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(abc9[2])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(abc9[3])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(abc9[4])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[5])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[8]) == 3:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,2).paragraphs[0].add_run(k9[2])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,3).paragraphs[0].add_run(abc9[0])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,4).paragraphs[0].add_run(abc9[1])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(abc9[2])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(abc9[3])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(abc9[4])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(abc9[5])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[6])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[8]) == 2:
        for i in klm2[8]:
            k9.append(i)

        table.cell(8,0).paragraphs[0].add_run(k9[0])
        table.cell(8,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(8,1).paragraphs[0].add_run(k9[1])
        table.cell(8,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,2).paragraphs[0].add_run(abc9[0])
        table.cell(8,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,3).paragraphs[0].add_run(abc9[1])
        table.cell(8,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,4).paragraphs[0].add_run(abc9[2])
        table.cell(8,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,5).paragraphs[0].add_run(abc9[3])
        table.cell(8,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,6).paragraphs[0].add_run(abc9[4])
        table.cell(8,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,7).paragraphs[0].add_run(abc9[5])
        table.cell(8,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,8).paragraphs[0].add_run(abc9[6])
        table.cell(8,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(8,9).paragraphs[0].add_run(abc9[7])
        table.cell(8,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[9]) == 10:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,3).paragraphs[0].add_run(k10[3])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,4).paragraphs[0].add_run(k10[4])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(k10[5])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(k10[6])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(k10[7])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(k10[8])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(k10[9])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[9]) == 9:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,3).paragraphs[0].add_run(k10[3])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,4).paragraphs[0].add_run(k10[4])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(k10[5])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(k10[6])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(k10[7])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(k10[8])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[0])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[9]) == 8:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,3).paragraphs[0].add_run(k10[3])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,4).paragraphs[0].add_run(k10[4])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(k10[5])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(k10[6])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(k10[7])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(abc10[0])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[1])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[9]) == 7:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,3).paragraphs[0].add_run(k10[3])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,4).paragraphs[0].add_run(k10[4])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(k10[5])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(k10[6])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(abc10[0])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(abc10[1])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[2])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[9]) == 6:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,3).paragraphs[0].add_run(k10[3])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,4).paragraphs[0].add_run(k10[4])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(k10[5])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(abc10[0])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(abc10[1])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(abc10[2])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[3])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[9]) == 5:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,3).paragraphs[0].add_run(k10[3])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,4).paragraphs[0].add_run(k10[4])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(abc10[0])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(abc10[1])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(abc10[2])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(abc10[3])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[4])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[9]) == 4:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,3).paragraphs[0].add_run(k10[3])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,4).paragraphs[0].add_run(abc10[0])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(abc10[1])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(abc10[2])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(abc10[3])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(abc10[4])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[5])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[9]) == 3:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,2).paragraphs[0].add_run(k10[2])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,3).paragraphs[0].add_run(abc10[0])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,4).paragraphs[0].add_run(abc10[1])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(abc10[2])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(abc10[3])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(abc10[4])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(abc10[5])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[6])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[9]) == 2:
        for i in klm2[9]:
            k10.append(i)

        table.cell(9,0).paragraphs[0].add_run(k10[0])
        table.cell(9,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(9,1).paragraphs[0].add_run(k10[1])
        table.cell(9,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,2).paragraphs[0].add_run(abc10[0])
        table.cell(9,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,3).paragraphs[0].add_run(abc10[1])
        table.cell(9,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,4).paragraphs[0].add_run(abc10[2])
        table.cell(9,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,5).paragraphs[0].add_run(abc10[3])
        table.cell(9,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,6).paragraphs[0].add_run(abc10[4])
        table.cell(9,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,7).paragraphs[0].add_run(abc10[5])
        table.cell(9,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,8).paragraphs[0].add_run(abc10[6])
        table.cell(9,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(9,9).paragraphs[0].add_run(abc10[7])
        table.cell(9,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[10]) == 10:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,3).paragraphs[0].add_run(k11[3])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,4).paragraphs[0].add_run(k11[4])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(k11[5])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(k11[6])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(k11[7])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(k11[8])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(k11[9])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[10]) == 9:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,3).paragraphs[0].add_run(k11[3])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,4).paragraphs[0].add_run(k11[4])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(k11[5])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(k11[6])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(k11[7])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(k11[8])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[0])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[10]) == 8:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,3).paragraphs[0].add_run(k11[3])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,4).paragraphs[0].add_run(k11[4])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(k11[5])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(k11[6])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(k11[7])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(abc11[0])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[1])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[10]) == 7:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,3).paragraphs[0].add_run(k11[3])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,4).paragraphs[0].add_run(k11[4])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(k11[5])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(k11[6])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(abc11[0])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(abc11[1])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[2])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[10]) == 6:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,3).paragraphs[0].add_run(k11[3])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,4).paragraphs[0].add_run(k11[4])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(k11[5])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(abc11[0])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(abc11[1])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(abc11[2])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[3])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[10]) == 5:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,3).paragraphs[0].add_run(k11[3])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,4).paragraphs[0].add_run(k11[4])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(abc11[0])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(abc11[1])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(abc11[2])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(abc11[3])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[4])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[10]) == 4:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,3).paragraphs[0].add_run(k11[3])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,4).paragraphs[0].add_run(abc11[0])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(abc11[1])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(abc11[2])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(abc11[3])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(abc11[4])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[5])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[10]) == 3:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,2).paragraphs[0].add_run(k11[2])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,3).paragraphs[0].add_run(abc11[0])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,4).paragraphs[0].add_run(abc11[1])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(abc11[2])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(abc11[3])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(abc11[4])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(abc11[5])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[6])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[10]) == 2:
        for i in klm2[10]:
            k11.append(i)

        table.cell(10,0).paragraphs[0].add_run(k11[0])
        table.cell(10,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(10,1).paragraphs[0].add_run(k11[1])
        table.cell(10,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,2).paragraphs[0].add_run(abc11[0])
        table.cell(10,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,3).paragraphs[0].add_run(abc11[1])
        table.cell(10,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,4).paragraphs[0].add_run(abc11[2])
        table.cell(10,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,5).paragraphs[0].add_run(abc11[3])
        table.cell(10,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,6).paragraphs[0].add_run(abc11[4])
        table.cell(10,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,7).paragraphs[0].add_run(abc11[5])
        table.cell(10,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,8).paragraphs[0].add_run(abc11[6])
        table.cell(10,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(10,9).paragraphs[0].add_run(abc11[7])
        table.cell(10,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[11]) == 10:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,3).paragraphs[0].add_run(k12[3])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,4).paragraphs[0].add_run(k12[4])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(k12[5])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(k12[6])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(k12[7])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(k12[8])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(k12[9])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[11]) == 9:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,3).paragraphs[0].add_run(k12[3])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,4).paragraphs[0].add_run(k12[4])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(k12[5])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(k12[6])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(k12[7])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(k12[8])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[0])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[11]) == 8:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,3).paragraphs[0].add_run(k12[3])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,4).paragraphs[0].add_run(k12[4])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(k12[5])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(k12[6])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(k12[7])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(abc12[0])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[1])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[11]) == 7:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,3).paragraphs[0].add_run(k12[3])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,4).paragraphs[0].add_run(k12[4])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(k12[5])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(k12[6])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(abc12[0])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(abc12[1])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[2])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[11]) == 6:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,3).paragraphs[0].add_run(k12[3])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,4).paragraphs[0].add_run(k12[4])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(k12[5])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(abc12[0])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(abc12[1])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(abc12[2])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[3])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[11]) == 5:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,3).paragraphs[0].add_run(k12[3])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,4).paragraphs[0].add_run(k12[4])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(abc12[0])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(abc12[1])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(abc12[2])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(abc12[3])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[4])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[11]) == 4:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,3).paragraphs[0].add_run(k12[3])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,4).paragraphs[0].add_run(abc12[0])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(abc12[1])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(abc12[2])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(abc12[3])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(abc12[4])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[5])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[11]) == 3:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,2).paragraphs[0].add_run(k12[2])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,3).paragraphs[0].add_run(abc12[0])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,4).paragraphs[0].add_run(abc12[1])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(abc12[2])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(abc12[3])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(abc12[4])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(abc12[5])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[6])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[11]) == 2:
        for i in klm2[11]:
            k12.append(i)

        table.cell(11,0).paragraphs[0].add_run(k12[0])
        table.cell(11,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(11,1).paragraphs[0].add_run(k12[1])
        table.cell(11,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,2).paragraphs[0].add_run(abc12[0])
        table.cell(11,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,3).paragraphs[0].add_run(abc12[1])
        table.cell(11,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,4).paragraphs[0].add_run(abc12[2])
        table.cell(11,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,5).paragraphs[0].add_run(abc12[3])
        table.cell(11,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,6).paragraphs[0].add_run(abc12[4])
        table.cell(11,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,7).paragraphs[0].add_run(abc12[5])
        table.cell(11,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,8).paragraphs[0].add_run(abc12[6])
        table.cell(11,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(11,9).paragraphs[0].add_run(abc12[7])
        table.cell(11,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[12]) == 10:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,3).paragraphs[0].add_run(k13[3])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,4).paragraphs[0].add_run(k13[4])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(k13[5])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(k13[6])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(k13[7])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(k13[8])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(k13[9])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[12]) == 9:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,3).paragraphs[0].add_run(k13[3])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,4).paragraphs[0].add_run(k13[4])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(k13[5])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(k13[6])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(k13[7])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(k13[8])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[0])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    if len(klm2[12]) == 8:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,3).paragraphs[0].add_run(k13[3])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,4).paragraphs[0].add_run(k13[4])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(k13[5])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(k13[6])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(k13[7])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(abc13[0])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[1])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[12]) == 7:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,3).paragraphs[0].add_run(k13[3])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,4).paragraphs[0].add_run(k13[4])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(k13[5])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(k13[6])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(abc13[0])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(abc13[1])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[2])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[12]) == 6:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,3).paragraphs[0].add_run(k13[3])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,4).paragraphs[0].add_run(k13[4])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(k13[5])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(abc13[0])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(abc13[1])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(abc13[2])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[3])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(klm2[12]) == 5:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,3).paragraphs[0].add_run(k13[3])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,4).paragraphs[0].add_run(k13[4])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(abc13[0])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(abc13[1])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(abc13[2])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(abc13[3])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[4])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[12]) == 4:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,3).paragraphs[0].add_run(k13[3])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,4).paragraphs[0].add_run(abc13[0])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(abc13[1])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(abc13[2])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(abc13[3])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(abc13[4])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[5])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[12]) == 3:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,2).paragraphs[0].add_run(k13[2])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,3).paragraphs[0].add_run(abc13[0])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,4).paragraphs[0].add_run(abc13[1])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(abc13[2])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(abc13[3])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(abc13[4])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(abc13[5])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[6])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if len(klm2[12]) == 2:
        for i in klm2[12]:
            k13.append(i)

        table.cell(12,0).paragraphs[0].add_run(k13[0])
        table.cell(12,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table.cell(12,1).paragraphs[0].add_run(k13[1])
        table.cell(12,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,2).paragraphs[0].add_run(abc13[0])
        table.cell(12,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,3).paragraphs[0].add_run(abc13[1])
        table.cell(12,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,4).paragraphs[0].add_run(abc13[2])
        table.cell(12,4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,5).paragraphs[0].add_run(abc13[3])
        table.cell(12,5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,6).paragraphs[0].add_run(abc13[4])
        table.cell(12,6).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,7).paragraphs[0].add_run(abc13[5])
        table.cell(12,7).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,8).paragraphs[0].add_run(abc13[6])
        table.cell(12,8).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(12,9).paragraphs[0].add_run(abc13[7])
        table.cell(12,9).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
  
    document.save("kelimeavi.docx")

    os.startfile("kelimeavi.docx")

buton=Button()
buton.config(text="KELİME AVI OLUŞTUR",command=kelime_avi,fg="white",bg="red",font="Times 16")
buton.place(relx = 0.20, rely = 0.30)

cikis = PhotoImage(file="cikis.png")

buton1=Button()
buton1.config(text="ÇIKIŞ",image=cikis,compound="top",command=pencere.destroy,width='50',fg="white",bg="red",font="Times 16")
buton1.place(relx = 0.82, rely = 0.60)

pencere.mainloop()
